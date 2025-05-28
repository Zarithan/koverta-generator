import streamlit as st
import pandas as pd
from docx import Document
import shutil
import os

# Load the postal code database
@st.cache_data
def ucitaj_bazu():
    return pd.read_csv("postanski_brojevi_rs.csv")

def format_postanski_broj(broj):
    return ' '.join(str(broj))

def popuni_prvu_stranu(template_path, output_path, ime, adresa, mesto, post_broj_fmt):
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    shutil.copy(template_path, output_path)
    doc = Document(output_path)
    for para in doc.paragraphs:
        if "ИМЕ И ПРЕЗИМЕ" in para.text or "НАЗИВ ОРГАНИЗАЦИЈЕ" in para.text:
            para.text = para.text.replace("ИМЕ И ПРЕЗИМЕ ИЛИ НАЗИВ ОРГАНИЗАЦИЈЕ", ime)
        elif "Адреса" in para.text:
            para.text = para.text.replace("Адреса", adresa)
        elif "Место" in para.text:
            para.text = para.text.replace("Место", mesto)
        elif "Поштански број" in para.text:
            para.text = para.text.replace("Поштански број", post_broj_fmt)
    doc.save(output_path)

def popuni_drugu_stranu(template_path, output_path, broj_predmeta, primalac, adresa):
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    shutil.copy(template_path, output_path)
    doc = Document(output_path)
    for para in doc.paragraphs:
        if "број предмета" in para.text:
            para.text = para.text.replace("број предмета", broj_predmeta)
        elif "Прималац" in para.text:
            para.text = para.text.replace("Прималац", primalac)
        elif "Адреса" in para.text:
            para.text = para.text.replace("Адреса", adresa)
    doc.save(output_path)

# Streamlit UI
st.title("📬 Генератор коверата (оба шаблона)")

ime = st.text_input("Прималац / Назив организације")
adresa = st.text_input("Адреса")
mesto = st.text_input("Место")
broj_predmeta = st.text_input("Број предмета")

if st.button("Генериши обе стране"):
    baza = ucitaj_bazu()
    rezultat = baza[baza['Место'].str.lower() == mesto.lower()]
    
    if rezultat.empty:
        st.error("⚠️ Место није пронађено у бази.")
    else:
        post_broj = int(rezultat.iloc[0]['Поштански број'])
        post_broj_fmt = format_postanski_broj(post_broj)

        os.makedirs("output", exist_ok=True)
        prva_out = f"output/prva_{ime.replace(' ', '_')}.docx"
        druga_out = f"output/druga_{ime.replace(' ', '_')}.docx"

        try:
            popuni_prvu_stranu("first_template.docx", prva_out, ime, adresa, mesto, post_broj_fmt)
            popuni_drugu_stranu("second_template.docx", druga_out, broj_predmeta, ime, adresa)

            with open(prva_out, "rb") as f1:
                st.download_button("⬇️ Преузми прву страну", f1, file_name=os.path.basename(prva_out))

            with open(druga_out, "rb") as f2:
                st.download_button("⬇️ Преузми другу страну", f2, file_name=os.path.basename(druga_out))
        except FileNotFoundError as e:
            st.error(str(e))
